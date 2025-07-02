import os
import logging
from typing import List, Dict, Any, Optional
from sqlalchemy.orm import Session
from models import Document
import aiofiles
import asyncio
from pathlib import Path
import pypdf
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Document processor for extracting text from uploaded files"""
    
    def __init__(self):
        self.chunk_size = int(os.getenv("CHUNK_SIZE", "1000"))
        self.chunk_overlap = int(os.getenv("CHUNK_OVERLAP", "200"))
    
    async def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text from different file types"""
        try:
            if file_type == "text/plain":
                return await self._extract_from_txt(file_path)
            elif file_type == "application/pdf":
                return await self._extract_from_pdf(file_path)
            elif "word" in file_type or file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._extract_from_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    async def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as f:
                    return await f.read()
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {str(e)}")
                return ""
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using pypdf"""
        try:
            # Run PDF processing in a thread to avoid blocking
            def process_pdf():
                try:
                    with open(file_path, 'rb') as file:
                        reader = pypdf.PdfReader(file)
                        text = ""
                        for page_num, page in enumerate(reader.pages):
                            try:
                                page_text = page.extract_text() or ""
                                text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                            except Exception as e:
                                logger.warning(f"Error extracting text from page {page_num + 1} in {file_path}: {str(e)}")
                                continue
                        return text.strip()
                except Exception as e:
                    logger.error(f"Error processing PDF {file_path}: {str(e)}")
                    return ""
            
            # Run in thread pool to avoid blocking the event loop
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, process_pdf)
            
            if not text:
                logger.warning(f"No text extracted from PDF: {file_path}")
                return f"[PDF file {Path(file_path).name} - No readable text content found]"
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF: {Path(file_path).name}")
            return text
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return f"[PDF file {Path(file_path).name} - Error processing file: {str(e)}]"
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx"""
        try:
            # Run DOCX processing in a thread to avoid blocking
            def process_docx():
                try:
                    doc = DocxDocument(file_path)
                    text_parts = []
                    
                    # Extract text from paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                    
                    return "\n\n".join(text_parts)
                    
                except Exception as e:
                    logger.error(f"Error processing DOCX {file_path}: {str(e)}")
                    return ""
            
            # Run in thread pool to avoid blocking the event loop
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, process_docx)
            
            if not text:
                logger.warning(f"No text extracted from DOCX: {file_path}")
                return f"[DOCX file {Path(file_path).name} - No readable text content found]"
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX: {Path(file_path).name}")
            return text
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return f"[DOCX file {Path(file_path).name} - Error processing file: {str(e)}]"
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks for better processing"""
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + self.chunk_size
            
            # Try to break at sentence or paragraph boundaries
            if end < text_length:
                # Look for sentence endings
                for i in range(end, max(start + self.chunk_size - 200, start), -1):
                    if text[i] in '.!?\n':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
            if start >= text_length:
                break
        
        return chunks

class RAGService:
    """RAG service for document retrieval and context generation"""
    
    def __init__(self):
        self.processor = DocumentProcessor()
        self.max_context_length = int(os.getenv("MAX_CONTEXT_LENGTH", "8000"))
        # Simple in-memory storage for document content (for demo purposes)
        # In production, use a proper vector database
        self.document_store: Dict[int, Dict[str, Any]] = {}
    
    async def process_document(self, document: Document, db: Session) -> bool:
        """Process a document and extract its content"""
        try:
            logger.info(f"Processing document: {document.fileName}")
            
            # Extract text from the document
            text_content = await self.processor.extract_text_from_file(
                document.filePath, 
                document.fileType
            )
            
            if not text_content:
                logger.warning(f"No text extracted from {document.fileName}")
                document.status = "failed"
                db.commit()
                return False
            
            # Chunk the text for better retrieval
            chunks = self.processor.chunk_text(text_content)
            
            # Store in our simple document store
            self.document_store[document.id] = {
                "fileName": document.fileName,
                "fileType": document.fileType,
                "userId": document.user_id,
                "fullText": text_content,
                "chunks": chunks,
                "uploadedBy": document.uploadedBy
            }
            
            logger.info(f"Document {document.fileName} processed into {len(chunks)} chunks")
            
            # Update document status
            document.status = "completed"
            db.commit()
            
            return True
            
        except Exception as e:
            logger.error(f"Error processing document {document.fileName}: {str(e)}")
            document.status = "failed"
            db.commit()
            return False
    
    async def retrieve_relevant_context(
        self, 
        query: str, 
        user_id: int, 
        db: Session
    ) -> tuple[str, List[Dict[str, Any]]]:
        """
        Retrieve relevant context from user's documents
        Returns: (context_text, source_documents)
        """
        try:
            # Get user's completed documents from database
            user_documents = db.query(Document).filter(
                Document.user_id == user_id,
                Document.status == "completed"
            ).all()
            
            if not user_documents:
                return "", []
            
            # For documents not in memory, process them
            for doc in user_documents:
                if doc.id not in self.document_store:
                    await self.process_document(doc, db)
            
            # Find relevant documents using keyword matching
            relevant_docs = []
            context_parts = []
            
            query_lower = query.lower()
            keywords = [word.strip() for word in query_lower.split() if len(word.strip()) > 2]
            
            for doc in user_documents:
                if doc.id not in self.document_store:
                    continue
                
                doc_data = self.document_store[doc.id]
                text_content = doc_data["fullText"]
                
                if not text_content:
                    continue
                
                text_lower = text_content.lower()
                
                # Calculate relevance score based on keyword matches
                relevance_score = 0
                for keyword in keywords:
                    relevance_score += text_lower.count(keyword)
                
                if relevance_score > 0:
                    # Extract relevant snippet
                    snippet = self._extract_relevant_snippet(text_content, query, keywords)
                    
                    relevant_docs.append({
                        "title": f"ONGC Document: {doc.fileName}",
                        "snippet": snippet,
                        "fileName": doc.fileName,
                        "fileUrl": f"/docs/{doc.id}/download",
                        "relevance": relevance_score
                    })
                    
                    context_parts.append(f"From {doc.fileName}:\n{snippet}")
            
            # Sort by relevance score
            relevant_docs.sort(key=lambda x: x["relevance"], reverse=True)
            
            # Take top 5 most relevant documents
            relevant_docs = relevant_docs[:5]
            context_parts = context_parts[:5]
            
            # Combine context parts, respecting max length
            context_text = "\n\n".join(context_parts)
            if len(context_text) > self.max_context_length:
                context_text = context_text[:self.max_context_length] + "..."
            
            logger.info(f"Retrieved context from {len(relevant_docs)} documents for query: {query[:50]}...")
            
            return context_text, relevant_docs
            
        except Exception as e:
            logger.error(f"Error retrieving context: {str(e)}")
            return "", []
    
    def _extract_relevant_snippet(self, text: str, query: str, keywords: List[str]) -> str:
        """Extract a relevant snippet from the text based on the query"""
        try:
            # Split text into sentences for better context
            sentences = []
            current_sentence = ""
            
            for char in text:
                current_sentence += char
                if char in '.!?\n' and len(current_sentence.strip()) > 10:
                    sentences.append(current_sentence.strip())
                    current_sentence = ""
            
            if current_sentence.strip():
                sentences.append(current_sentence.strip())
            
            # Find sentences with the most keyword matches
            best_sentences = []
            for sentence in sentences:
                sentence_lower = sentence.lower()
                matches = sum(1 for keyword in keywords if keyword in sentence_lower)
                if matches > 0:
                    best_sentences.append((sentence, matches))
            
            # Sort by number of matches and take top sentences
            best_sentences.sort(key=lambda x: x[1], reverse=True)
            
            if best_sentences:
                # Combine top matching sentences
                snippet_parts = []
                total_length = 0
                
                for sentence, _ in best_sentences[:3]:  # Take top 3 sentences
                    if total_length + len(sentence) < 500:
                        snippet_parts.append(sentence)
                        total_length += len(sentence)
                    else:
                        break
                
                snippet = " ".join(snippet_parts)
                if len(snippet) > 500:
                    snippet = snippet[:500] + "..."
                
                return snippet
            
            # Fallback: return first 300 characters
            return text[:300] + "..." if len(text) > 300 else text
            
        except Exception as e:
            logger.error(f"Error extracting snippet: {str(e)}")
            return text[:200] + "..." if len(text) > 200 else text

# Global service instance
_rag_service: Optional[RAGService] = None

def get_rag_service() -> RAGService:
    """Get or create the global RAG service instance"""
    global _rag_service
    if _rag_service is None:
        _rag_service = RAGService()
    return _rag_service